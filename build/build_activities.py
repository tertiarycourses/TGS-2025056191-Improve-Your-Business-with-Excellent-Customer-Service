#!/usr/bin/env python3
"""Per-activity handouts — Improve Your Business with Excellent Customer Service.

Writes ONE FOLDER PER ACTIVITY:

    courseware/activities/
        README.md                       index of all nine activities
        activity-01-.../
            activity-01-....md                        the brief (scenario · questions · steps · debrief)
            activity-01-....pdf                       the same brief rendered to PDF (handout)
            activity-01-...-observer-checklist.pdf    role plays only — scored observation sheet
            activity-01-...-reflection.pdf            every activity — post-debrief reflection
        activity-02-.../
        ...

Content comes from data_domain1/2 — the same single source that drives the
slide deck, the Lesson Plan and the Learner Guide, so the four can never drift.

Markdown -> PDF is done with LibreOffice via an intermediate DOCX built with
python-docx, so the PDF carries the house fonts and a proper cover block
(pandoc/LaTeX are not assumed to be installed).
"""
import os, re, sys, subprocess, shutil

HERE = os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0, HERE)
import course_data as C
from data_domain1 import DOMAIN1
from data_domain2 import DOMAIN2
ACT = DOMAIN1 + DOMAIN2
import prodoc
from data_checklists import CHECKLISTS, REFLECTION_PROMPT
from docx import Document
from docx.shared import Pt, RGBColor, Inches

REPO = os.path.dirname(HERE)                       # .../courseware
ASSETS = os.path.join(REPO, "assets")
OUT = os.path.join(REPO, "activities")

SOFFICE = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
BRAND = RGBColor(0x1F, 0x6F, 0xEB); GREEN = RGBColor(0x16, 0x84, 0x5B)
GREY = RGBColor(0x55, 0x5B, 0x66)

TOPICS = {t["num"]: t for t in C.TOPICS}


def slug(s):
    s = re.sub(r"[^a-zA-Z0-9]+", "-", s.lower()).strip("-")
    return re.sub(r"-+", "-", s)[:60]


# ---------------------------------------------------------------- markdown
def render_md(a):
    t = TOPICS[a["topic"]]
    L = []
    L.append(f"# Activity {a['num']} — {a['title']}")
    L.append("")
    L.append(f"**Course:** {C.TITLE} ({C.COURSE_CODE})  ")
    L.append(f"**Topic {t['num']}:** {t['title']}  ")
    L.append(f"**Mapping:** {a['objective']}  ")
    L.append(f"**Format:** {a['type']}  ·  **Duration:** {a['minutes']} minutes  ")
    L.append(f"**Roles:** {a['roles']}")
    L.append("")
    L.append("---")
    L.append("")
    L.append("## The Scenario")
    L.append("")
    for para in [x.strip() for x in a["scenario"].split("\n\n") if x.strip()]:
        L.append(para)
        L.append("")
    L.append("## What You Will Produce")
    L.append("")
    L.append(a["build"])
    L.append("")
    L.append(f"*Materials: {a['services']}.*")
    L.append("")
    L.append("## Discussion Questions")
    L.append("")
    for i, q in enumerate(a["questions"], 1):
        L.append(f"{i}. {q}")
    L.append("")
    L.append("## Step-by-Step")
    L.append("")
    for i, (instr, _cmd) in enumerate(a["steps"], 1):
        L.append(f"{i}. {instr}")
    L.append("")
    L.append("## Debrief — What a Strong Answer Looks Like")
    L.append("")
    for d in a["debrief"]:
        L.append(f"- {d}")
    L.append("")
    L.append("## Check Your Work")
    L.append("")
    L.append(a["test"])
    L.append("")
    L.append("---")
    L.append("")
    L.append(f"*{C.ORG} · {C.UEN} · {C.TITLE} ({C.COURSE_CODE}) · Version {C.VERSION}*")
    L.append("")
    return "\n".join(L)


# ---------------------------------------------------------------- docx -> pdf
def render_docx(a, path):
    t = TOPICS[a["topic"]]
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    prodoc.style_headings(doc)

    # header block
    p = doc.add_paragraph()
    r = p.add_run(f"{C.ORG}   ·   {C.UEN}")
    r.font.size = Pt(9); r.font.color.rgb = GREY
    p = doc.add_paragraph()
    r = p.add_run(f"{C.TITLE}   ·   {C.COURSE_CODE}")
    r.font.size = Pt(9); r.font.color.rgb = GREY

    doc.add_heading(f"Activity {a['num']} — {a['title']}", level=1)

    meta = doc.add_table(rows=0, cols=2); meta.style = "Table Grid"
    for k, v in [("Topic", f"{t['num']} — {t['title']}"),
                 ("Mapping", a["objective"]),
                 ("Format", a["type"]),
                 ("Duration", f"{a['minutes']} minutes"),
                 ("Roles", a["roles"])]:
        c = meta.add_row().cells
        c[0].text = ""; rr = c[0].paragraphs[0].add_run(k); rr.bold = True; rr.font.size = Pt(9.5)
        prodoc._shade_cell(c[0], "E8F0FE")
        c[1].text = ""; c[1].paragraphs[0].add_run(v).font.size = Pt(9.5)

    def head(txt, color=BRAND):
        pp = doc.add_paragraph()
        rr = pp.add_run(txt); rr.bold = True; rr.font.size = Pt(12); rr.font.color.rgb = color

    head("The Scenario")
    for para in [x.strip() for x in a["scenario"].split("\n\n") if x.strip()]:
        doc.add_paragraph(para)

    head("What You Will Produce")
    doc.add_paragraph(a["build"])
    pp = doc.add_paragraph(); rr = pp.add_run(f"Materials: {a['services']}.")
    rr.italic = True; rr.font.size = Pt(10); rr.font.color.rgb = GREY

    def numbered(items):
        """Explicit numbering — Word's List Number style continues across
        separate lists, which would make Step-by-Step start at 7."""
        for i, itxt in enumerate(items, 1):
            pp = doc.add_paragraph()
            pp.paragraph_format.left_indent = Pt(18)
            pp.paragraph_format.first_line_indent = Pt(-18)
            rr = pp.add_run(f"{i}.  "); rr.bold = True
            pp.add_run(itxt)

    head("Discussion Questions")
    numbered(a["questions"])

    head("Step-by-Step")
    numbered([instr for instr, _cmd in a["steps"]])

    head("Debrief — What a Strong Answer Looks Like", GREEN)
    for d in a["debrief"]:
        doc.add_paragraph(d, style="List Bullet")

    head("Check Your Work", GREEN)
    doc.add_paragraph(a["test"])

    prodoc.add_page_numbers(doc)
    doc.save(path)


def to_pdf(docx_path, outdir):
    subprocess.run([SOFFICE, "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
                   check=True, capture_output=True)


# ------------------------------------------------- role-play observer checklist
def _doc_header(doc, a, kicker):
    """Shared header block for the checklist / reflection sheets."""
    p = doc.add_paragraph()
    r = p.add_run(f"{C.ORG}   ·   {C.UEN}")
    r.font.size = Pt(9); r.font.color.rgb = GREY
    p = doc.add_paragraph()
    r = p.add_run(f"{C.TITLE}   ·   {C.COURSE_CODE}")
    r.font.size = Pt(9); r.font.color.rgb = GREY
    doc.add_heading(f"Activity {a['num']} — {kicker}", level=1)
    p = doc.add_paragraph()
    r = p.add_run(a["title"]); r.bold = True; r.font.size = Pt(11.5)


def _rule_row(cell, text, bold=False, size=9.5, fill=None, color=None):
    cell.text = ""
    rr = cell.paragraphs[0].add_run(text)
    rr.bold = bold; rr.font.size = Pt(size); rr.font.name = "Arial"
    if color: rr.font.color.rgb = color
    if fill: prodoc._shade_cell(cell, fill)


def render_checklist(a, cl, path):
    """Observer checklist for a role-play activity — one scored row per criterion."""
    doc = Document()
    for sec in doc.sections:              # narrower margins: the table needs the width
        sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    prodoc.style_headings(doc)
    _doc_header(doc, a, "Role-Play Observer Checklist")

    meta = doc.add_table(rows=0, cols=2); meta.style = "Table Grid"
    for k, v in [("Focus of this observation", cl["focus"]),
                 ("You are observing", cl["observing"]),
                 ("Setup", a["roles"]),
                 ("Duration", f"{a['minutes']} minutes")]:
        c = meta.add_row().cells
        _rule_row(c[0], k, bold=True, fill="E8F0FE")
        _rule_row(c[1], v)

    p = doc.add_paragraph()
    r = p.add_run("Observer: ______________________     Performer: ______________________     "
                  "Round:  1  /  2")
    r.font.size = Pt(10)

    p = doc.add_paragraph()
    r = p.add_run("Tick DONE only where you actually saw the behaviour. "
                  "'What good looks like' is the standard — anything less is NOT YET.")
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY

    t = doc.add_table(rows=0, cols=5); t.style = "Table Grid"
    hdr = t.add_row().cells
    for i, h in enumerate(["Code", "What to observe", "What good looks like", "Done", "Not yet"]):
        _rule_row(hdr[i], h, bold=True, size=9,
                  color=RGBColor(0xFF, 0xFF, 0xFF), fill="1F6FEB")
    for code, crit, good in cl["rows"]:
        c = t.add_row().cells
        _rule_row(c[0], code, bold=True, size=8.5, fill="E8F0FE")
        _rule_row(c[1], crit, bold=True, size=8.5)
        _rule_row(c[2], good, size=8.5)
        _rule_row(c[3], "", size=9)
        _rule_row(c[4], "", size=9)
    # Fix the column widths: without autofit off, Word/LibreOffice re-flow the
    # table from content and the "what good looks like" column collapses.
    t.autofit = False
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _El
    tblPr = t._tbl.tblPr
    lay = _El("w:tblLayout"); lay.set(_qn("w:type"), "fixed"); tblPr.append(lay)
    # dxa = twentieths of a point. Total 7.1in of usable width at 0.7in margins.
    widths_dxa = [660, 1900, 4420, 680, 680]
    grid = _El("w:tblGrid")
    for w in widths_dxa:
        gc = _El("w:gridCol"); gc.set(_qn("w:w"), str(w)); grid.append(gc)
    t._tbl.insert(1, grid)
    for row in t.rows:
        for i, w in enumerate(widths_dxa):
            tc = row.cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(_qn("w:tcW")):
                tcPr.remove(old)
            tcW = _El("w:tcW"); tcW.set(_qn("w:w"), str(w))
            tcW.set(_qn("w:type"), "dxa"); tcPr.append(tcW)

    def head(txt, color=BRAND):
        pp = doc.add_paragraph()
        rr = pp.add_run(txt); rr.bold = True; rr.font.size = Pt(12); rr.font.color.rgb = color

    for prompt in ("Strongest moment — what worked, and why",
                   "One thing to change in the next round",
                   "The performer's own view (ask them first, before you give feedback)"):
        pp = doc.add_paragraph()
        pp.paragraph_format.space_before = Pt(6); pp.paragraph_format.space_after = Pt(2)
        rr = pp.add_run(prompt); rr.bold = True; rr.font.size = Pt(10)
        rr.font.color.rgb = BRAND
        ln = doc.add_paragraph()
        ln.paragraph_format.space_after = Pt(2)
        lr = ln.add_run("_" * 104)
        lr.font.size = Pt(9); lr.font.color.rgb = RGBColor(0xB8, 0xC2, 0xCE)

    pp = doc.add_paragraph()
    pp.paragraph_format.space_before = Pt(4)
    rr = pp.add_run("Check your work:  " + a["test"])
    rr.italic = True; rr.font.size = Pt(8.5); rr.font.color.rgb = GREY

    prodoc.add_page_numbers(doc)
    doc.save(path)


def _lines(doc, n=3):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run("_" * 96)
        r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xB8, 0xC2, 0xCE)


def render_reflection(a, prompt, path):
    """Post-activity reflection worksheet — transfers the activity to the learner's own workplace."""
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    prodoc.style_headings(doc)
    _doc_header(doc, a, "Reflection Worksheet")

    p = doc.add_paragraph()
    r = p.add_run("Name: ______________________________     Date: ____________________")
    r.font.size = Pt(10)

    p = doc.add_paragraph()
    r = p.add_run("Complete this after the debrief. It is not assessed — it is how you carry the "
                  "activity back to your own workplace. Keep it with your Learner Guide.")
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY

    def q(txt, lines=3, color=BRAND):
        pp = doc.add_paragraph()
        rr = pp.add_run(txt); rr.bold = True; rr.font.size = Pt(11); rr.font.color.rgb = color
        _lines(doc, lines)

    q("1.  What did you do in this activity that you would do again?")
    q("2.  What did you see someone else do that you want to steal?")
    q("3.  Where did you get it wrong — and what did that cost the customer?")
    q("4.  " + prompt, 3)
    q("5.  One specific thing you will do differently at work next week:", 2)
    q("6.  What would have to be true for that change to actually stick?", 2)

    rect = doc.add_table(rows=1, cols=1); rect.style = "Table Grid"
    cell = rect.rows[0].cells[0]
    cell.text = ""
    rr = cell.paragraphs[0].add_run("Links to the assessment")
    rr.bold = True; rr.font.size = Pt(10); rr.font.color.rgb = GREEN
    b = cell.add_paragraph()
    br = b.add_run(a["objective"]); br.font.size = Pt(9.5)
    b2 = cell.add_paragraph()
    b2r = b2.add_run("Check your work:  " + a["test"]); b2r.font.size = Pt(9.5)

    prodoc.add_page_numbers(doc)
    doc.save(path)


# ---------------------------------------------------------------- build
def main():
    os.makedirs(OUT, exist_ok=True)
    index = ["# Activities — " + C.TITLE, "",
             f"**Course code:** {C.COURSE_CODE}  ·  **Version:** {C.VERSION}  ·  {C.ORG}", "",
             "Nine real-world Singapore case-study activities. Each folder contains the activity "
             "brief (Markdown + printable PDF), a **reflection worksheet**, and — for the six "
             "role-play activities — a **role-play observer checklist**.", "",
             "| # | Activity | Topic | Format | Duration | Checklist | Folder |",
             "|---|---|---|---|---|---|---|"]

    for a in ACT:
        t = TOPICS[a["topic"]]
        folder_name = f"activity-{a['num']:02d}-{slug(a['title'])}"
        folder = os.path.join(OUT, folder_name)
        os.makedirs(folder, exist_ok=True)
        base = folder_name

        md_path = os.path.join(folder, base + ".md")
        with open(md_path, "w") as f:
            f.write(render_md(a))

        docx_path = os.path.join(folder, base + ".docx")
        render_docx(a, docx_path)
        to_pdf(docx_path, folder)
        os.remove(docx_path)          # DOCX was only the PDF intermediate

        made = ["brief"]

        # Role-play observer checklist — only for the role-play activities.
        cl = CHECKLISTS.get(a["num"])
        if cl:
            cbase = f"{base}-observer-checklist"
            cdocx = os.path.join(folder, cbase + ".docx")
            render_checklist(a, cl, cdocx)
            to_pdf(cdocx, folder)
            os.remove(cdocx)
            made.append("checklist")

        # Reflection worksheet — every activity.
        rbase = f"{base}-reflection"
        rdocx = os.path.join(folder, rbase + ".docx")
        render_reflection(a, REFLECTION_PROMPT[a["num"]], rdocx)
        to_pdf(rdocx, folder)
        os.remove(rdocx)
        made.append("reflection")

        pdf_path = os.path.join(folder, base + ".pdf")
        ok = os.path.exists(pdf_path)
        print(f"  {'✓' if ok else '✗'} {folder_name}/  ({' + '.join(made)})")

        index.append(f"| {a['num']} | {a['title']} | T{t['num']} {t['title']} | {a['type']} | "
                     f"{a['minutes']} min | {'✅' if cl else '—'} | [`{folder_name}/`]({folder_name}/) |")

    index += ["", "## What is in each folder", "",
              "| File | Who uses it | When |",
              "|---|---|---|",
              "| `*-<activity>.md` / `.pdf` | Learner | During the activity — scenario, questions, step-by-step |",
              "| `*-observer-checklist.pdf` | The observing learner | During the role play — tick each behaviour as it happens |",
              "| `*-reflection.pdf` | Learner | After the debrief — transfer to their own workplace |",
              "",
              "The observer checklist criteria are taken from that activity's own debrief points, so "
              "what the observer scores is exactly what the trainer debriefs and what the assessor "
              "looks for in the Role Play (RP) assessment.",
              "", "## How the activities map to the assessment", "",
              "| Activity | Abilities practised | Assessment instrument |",
              "|---|---|---|"]
    for a in ACT:
        codes = [c for c in ("A1", "A2", "A3") if c in a["objective"]]
        ks = [c for c in ("K1", "K2", "K3") if c in a["objective"]]
        instr = "Role Play (RP)" if codes else "Oral Questioning (OQ)"
        if codes and ks:
            instr = "Role Play (RP) + Oral Questioning (OQ)"
        index.append(f"| {a['num']}. {a['title']} | {', '.join(codes + ks) or '—'} | {instr} |")
    index.append("")

    with open(os.path.join(OUT, "README.md"), "w") as f:
        f.write("\n".join(index))
    print("Saved", os.path.join(OUT, "README.md"))


if __name__ == "__main__":
    main()
