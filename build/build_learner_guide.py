#!/usr/bin/env python3
"""Learner Guide (LG) DOCX — Improve Your Business with Excellent Customer Service.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per topic (concepts explained in prose) and one section per
activity (scenario · discussion questions · DETAILED step-by-step · debrief ·
test). All content is driven by course_data + data_domain1/2, keeping the LG
100% aligned with the slide deck, Lesson Plan and the activities folder.

HARD RULE (wsq-learner-guide): deliverables are the DOCX and its PDF ONLY —
no Markdown mirror is kept in the repo. Markdown is emitted only as a
throwaway intermediate for the per-activity handouts, which are built by
build_activities.py, not here.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0, HERE)
import course_data as C
from data_domain1 import DOMAIN1
from data_domain2 import DOMAIN2
ACT = DOMAIN1 + DOMAIN2
import prodoc
REPO = os.path.dirname(HERE)                    # .../courseware
ASSETS = os.path.join(REPO, "assets")
OUTDIR = os.path.join(REPO, "courseware")

# ---------------- block DSL ----------------
B = []
def h1(t): B.append(("h1", t))
def h2(t): B.append(("h2", t))
def h3(t): B.append(("h3", t))
def p(t):  B.append(("p", t))
def bullets(xs): B.append(("bullets", xs))
def numbered(xs): B.append(("numbered", xs))
def steps(xs): B.append(("steps", xs))
def note(t): B.append(("note", t))
def rule(): B.append(("rule",))
def dl(pairs): B.append(("dl", pairs))

# ================================================================= content
h1("Introduction")
p(f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by "
  f"{C.ORG}. It covers all nine topics across the two learning units, and gives the full "
  "step-by-step instructions for all nine case-study activities you will complete in class.")
p(f"The course is mapped to the Skills Framework TSC {C.TSC_TITLE} ({C.TSC_CODE}). Everything in "
  "this guide is assessable: the Oral Questioning paper draws on the knowledge (K1, K2, K3) "
  "explained in the topic sections, and the Role Play draws on the abilities (A1, A2, A3) you "
  "practise in the activities.")

h1("How to Use This Guide")
bullets([
 "Read the topic section before its activity — the concepts are what the activity asks you to apply.",
 "During each activity, work from the Step-by-step section. The slides show only the scenario and "
 "the discussion questions; the detailed procedure lives here.",
 "After each activity, read the Debrief section. It states what a strong answer looks like, which is "
 "also what the assessor is looking for.",
 "The 'Check your work' line at the end of every activity is your self-assessment standard.",
 "This guide is an open-book reference in the final assessment — bring it with you.",
 "A printable one-page brief for each activity is in the activities folder (activities/activity-NN/).",
])

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h1("Skills Framework Alignment")
p(f"TSC: {C.TSC_TITLE}  ·  TSC Code: {C.TSC_CODE}")
h3("Abilities assessed by the Role Play (RP)")
bullets([f"{c} — {d}" for c, d in C.TSC_ABILITIES])
h3("Knowledge assessed by Oral Questioning (OQ)")
bullets([f"{c} — {d}" for c, d in C.TSC_KNOWLEDGE])

h1("Assessment Overview")
bullets([C.ASSESSMENT["written"], C.ASSESSMENT["practical"],
         "Format: Open Book — these course slides, this Learner Guide and approved materials only.",
         "Oral Clarification of up to 10 minutes may be conducted 1:1 to close minor performance gaps. "
         "This is not counted in the assessment duration.",
         C.ASSESSMENT["note"],
         "Submit your completed answers on the LMS at https://lms-tms.tertiaryinfotech.com/."])

# ---------------- per-topic, per-activity ----------------
for t in C.TOPICS:
    h1(f"Topic {t['num']} — {t['title']}")
    p(t["subtitle"])
    p(f"Learning unit: LU{t['lu']}  ·  Mapping: {t['weighting']}")
    h3("Key concepts")
    for title, caption in t["concepts"]:
        B.append(("concept", title, caption))

    for a in [x for x in ACT if x["topic"] == t["num"]]:
        h2(f"Activity {a['num']} — {a['title']}")
        p(f"Mapping: {a['objective']}")
        p(f"Format: {a['type']}  ·  Duration: {a['minutes']} minutes  ·  {a['roles']}")

        h3("Scenario")
        for para in [x.strip() for x in a["scenario"].split("\n\n") if x.strip()]:
            p(para)

        h3("What you will produce")
        p(a["build"] + f"   (Materials: {a['services']}.)")

        h3("Discussion questions")
        numbered(a["questions"])

        h3("Step-by-step")
        steps([instr for instr, _cmd in a["steps"]])

        h3("Debrief — what a strong answer looks like")
        bullets(a["debrief"])

        h3("Check your work")
        p(a["test"])
        note(f"A printable one-page brief for this activity is in activities/activity-{a['num']:02d}/.")
        rule()

# ---------------- closing sections ----------------
h1("Preparing for the Assessment")
h3("Oral Questioning (OQ) — 30 minutes, 3 questions")
p("The OQ tests your KNOWLEDGE. Each question maps to one knowledge factor. Prepare by being able "
  "to explain, in your own words and with an example:")
bullets([
 "K1 — Principles of effective communication: active listening, paraphrasing, open vs closed "
 "questioning, positive language, body language, tone, and channel-appropriate written style. "
 "(Topics 1–6.)",
 "K2 — Customer feedback channels: surveys (CSAT, NPS, CES), online reviews, support tickets, "
 "social listening, direct interviews — with one advantage and one limitation of each. (Topic 7.)",
 "K3 — Operation and process personnel feedback channels: shift huddles, incident logs, team "
 "retrospectives, internal escalation reports, QA scorecards and defect channels — and how they "
 "convert customer complaints into service improvements. (Topics 7 and 9.)",
])
h3("Role Play (RP) — 30 minutes, 2 role plays")
p("The RP tests your ABILITY. The assessor plays the customer. Prepare by rehearsing the routines "
  "you practised in the activities:")
bullets([
 "A1 — Carry out collection of customer feedback: greet, build rapport, use at least two feedback "
 "collection methods in the conversation, and record what you heard. (Activities 1, 4, 5, 7.)",
 "A2 — Determine customer needs and expectations: open questions first, paraphrase to confirm, "
 "separate what was said from what is needed, then map to the four levels. (Activities 3, 4.)",
 "A3 — Determine areas of improvement from feedback: identify the root cause behind the complaint, "
 "recommend a specific improvement with an owner, and close the loop with the customer. "
 "(Activities 7, 8, 9.)",
])
h3("Common reasons candidates are marked Not Yet Competent")
bullets([
 "Jumping to a solution before diagnosing the need — always paraphrase and confirm first.",
 "Apologising conditionally ('if you were inconvenienced') instead of unconditionally.",
 "Recommending a vague improvement ('improve communication') with no owner, date or measure.",
 "Failing to name a feedback channel when asked how feedback would be collected.",
 "Continuing to serve after abuse, instead of warning and escalating.",
 "Ending the interaction without stating concrete next steps, a time and an owner.",
])

h1("Quick Reference — The Frameworks in This Course")
dl([
 ("The service chain", "Internal support → frontline delivery → customer experience → loyalty and revenue. "
  "A break anywhere upstream surfaces as a frontline failure."),
 ("Four principles of good service", "Personalised · Competent · Convenient · Proactive. All four must hold."),
 ("Active listening", "Attend → Absorb → Paraphrase → Confirm → Act. Paraphrasing is what makes listening visible."),
 ("Four levels of addressing needs", "1 Understand the problem · 2 Meet the basic need · 3 Think outside the "
  "box · 4 Go the extra mile. You cannot skip a rung."),
 ("Park and return", "Note where you stopped so the customer never has to repeat themselves."),
 ("The six-stage call", "Greet → Listen → Paraphrase → Hold → Resolve → Close."),
 ("The four-part hold", "Ask permission · state the reason · give a duration · thank them on return."),
 ("Five-part service email", "Greeting → Acknowledge → Solution → Next steps → Sign-off."),
 ("Channel selection", "Match the channel to the issue: urgency on one axis, complexity on the other."),
 ("CSAT / NPS / CES", "CSAT rates a moment · NPS rates the relationship · CES rates the customer's effort."),
 ("The closed loop", "Collect → Analyse → Act → Tell them. Only step 4 is visible to the customer."),
 ("HEARD", "Hear · Empathize · Apologize · Resolve · Diagnose."),
 ("LAST", "Listen · Apologize · Solve · Thank."),
 ("Perceived justice", "Customers judge the outcome, the process AND how they were treated."),
 ("Functional vs hierarchical escalation", "Sideways for expertise; upward for authority."),
 ("The three-strike rule", "Two calm behavioural warnings, then escalate or end. Abuse skips the warnings."),
])

h1("Glossary")
dl([
 ("Customer service", "The whole of the support a customer receives before, during and after a purchase."),
 ("External customer", "An individual or business that buys your products or services."),
 ("Internal customer", "A colleague or department that depends on your work output."),
 ("Service recovery", "Resolving a service failure so as to restore the customer's trust, not merely fix the fault."),
 ("Active listening", "Attending fully, paraphrasing back and confirming before responding."),
 ("Positive language", "Framing a response around what IS possible rather than what is not."),
 ("Netiquette", "The conventions of polite, professional written communication online."),
 ("CSAT", "Customer Satisfaction score — how satisfied a customer was with a specific interaction."),
 ("NPS", "Net Promoter Score — how likely a customer is to recommend you."),
 ("CES", "Customer Effort Score — how much effort the customer had to spend to get resolution."),
 ("Closing the loop", "Telling the customer what changed as a result of the feedback they gave."),
 ("Root cause", "The underlying reason a problem occurs, as opposed to the symptom the customer reports."),
 ("De-escalation", "Reducing a customer's emotional intensity so the problem can be addressed."),
 ("Escalation", "Transferring an issue to someone with the required expertise or authority."),
 ("POHA", "Protection from Harassment Act 2014 — Singapore legislation criminalising harassment, "
  "including abuse of workers at work."),
 ("TAFEP", "Tripartite Alliance for Fair & Progressive Employment Practices — receives workplace "
  "harassment reports in Singapore."),
 ("TRAQOM", "The SSG system used for digital attendance and course feedback on WSQ-funded courses."),
])

# ================================================================= render DOCX
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27)
GREY = RGBColor(0x55, 0x5B, 0x66); GREEN = RGBColor(0x16, 0x84, 0x5B)

doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc, "LEARNER GUIDE", C.TITLE, C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc, [
    ("1.0", "1 March 2026",
     "Initial release — 9 topics across 2 learning units, aligned to TSC EPW-CEX-3034-1.1.", C.TRAINER),
    (C.VERSION.lstrip("v"), C.VERSION_DATE,
     "Major content revision. Nine topics deepened from current industry practice (Zendesk, Help Scout, "
     "Qualtrics, SurveyMonkey, Coursera, Tidio, Intercom, Indeed SG, Global Response). Nine real-world "
     "Singapore case-study activities added, each with scenario, discussion questions, detailed "
     "step-by-step, trainer debrief and a self-check standard. Added assessment preparation and "
     "framework quick-reference sections.", C.TRAINER),
])
prodoc.add_toc(doc)

for kind, *rest in B:
    if kind == "h1":
        doc.add_heading(rest[0], level=1)
    elif kind == "h2":
        doc.add_heading(rest[0], level=2)
    elif kind == "h3":
        para = doc.add_paragraph()
        r = para.add_run(rest[0]); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BRAND
    elif kind == "p":
        doc.add_paragraph(rest[0])
    elif kind == "concept":
        para = doc.add_paragraph(style="List Bullet")
        r = para.add_run(rest[0] + " — "); r.bold = True
        para.add_run(rest[1])
    elif kind == "bullets":
        for x in rest[0]:
            doc.add_paragraph(x, style="List Bullet")
    elif kind in ("numbered", "steps"):
        # Explicit numbering: Word's List Number style continues across separate
        # lists, so the Step-by-step list would carry on from the questions.
        for i, x in enumerate(rest[0], 1):
            pp = doc.add_paragraph()
            pp.paragraph_format.left_indent = Pt(20)
            pp.paragraph_format.first_line_indent = Pt(-20)
            pp.paragraph_format.space_after = Pt(3)
            rr = pp.add_run(f"{i}.  "); rr.bold = True
            pp.add_run(x)
    elif kind == "note":
        para = doc.add_paragraph()
        r = para.add_run("Note: "); r.bold = True; r.font.color.rgb = BRAND
        para.add_run(rest[0]).font.size = Pt(10)
    elif kind == "rule":
        doc.add_paragraph("")
    elif kind == "dl":
        for term, defn in rest[0]:
            para = doc.add_paragraph(style="List Bullet")
            r = para.add_run(term + " — "); r.bold = True
            para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
os.makedirs(OUTDIR, exist_ok=True)
DOCX_OUT = os.path.join(OUTDIR, f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved", DOCX_OUT)
