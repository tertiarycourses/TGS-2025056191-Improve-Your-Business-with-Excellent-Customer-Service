#!/usr/bin/env python3
"""Lesson Plan (LP) DOCX — Improve Your Business with Excellent Customer Service.

House format: cover page + Document Version Control Record + auto TOC +
Arial 11pt body + colour-coded day schedule table.

Course design (per the approved Course Proposal TPG-2025092324):
  7 h classroom facilitation + 1 h assessment (OQ 30 min + RP 30 min) = 8 h.
Topics and activities come from course_data + data_domain1/2 so the LP stays
aligned with the deck, the Learner Guide and the activities folder.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0, HERE)
import course_data as C
from data_domain1 import DOMAIN1
from data_domain2 import DOMAIN2
ACT = DOMAIN1 + DOMAIN2
import prodoc
REPO = os.path.dirname(HERE)                      # .../courseware
ASSETS = os.path.join(REPO, "assets")
OUTDIR = os.path.join(REPO, "courseware")

BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27)
GREY = RGBColor(0x55, 0x5B, 0x66)
HEADER_FILL = "1F6FEB"; TOPIC_FILL = "E8F0FE"; BREAK_FILL = "FFF4E5"
LUNCH_FILL = "FDE9D9"; ASSESS_FILL = "E8F7EE"


def act(n):
    a = [x for x in ACT if x["num"] == n][0]
    return f"Activity {n}: {a['title']} ({a['minutes']} min, {a['type']})"


# ---------------------------------------------------------------- schedule
# 9:00am-6:00pm. Exactly 420 min classroom facilitation + 60 min assessment,
# plus a 1-hour lunch (not counted). Tea breaks are counted within training time.
TOPIC_TEXT = {
 0: "Topic 1 — Understanding Customers & Customer Service: internal vs external customers, the service "
    "chain, the four principles of good service, the business case for service (K1)",
 1: "Topic 2 — Establishing Your Service Attitude: first impressions, the four pillars, positive language, "
    "energy management and resilience (K1)",
 2: "Topic 3 — Identifying and Addressing Customer Needs: active listening, open vs closed questioning, the "
    "four levels of addressing needs, said vs needed (K1, A2)",
 3: "Topic 4 — In-Person Customer Service: strengths and limits of face-to-face, body language, handling "
    "walk-in and at-your-desk requests, the park-and-return method (K1, A1)",
 4: "Topic 5 — Customer Service Over the Phone: telephone etiquette, the six-stage call structure, the "
    "four-part hold, paraphrasing to confirm (K1, A1)",
 5: "Topic 6 — Customer Service via Email and Chat: netiquette, the five-part email structure, written "
    "tone, choosing between email and live chat (K1, K2, A1)",
 6: "Topic 7 — Generating Return Business from Feedback: feedback channels, CSAT/NPS/CES, the silent "
    "majority, closing the loop, operational and personnel feedback channels (K2, K3, A1, A3)",
 7: "Topic 8 — Recovering Difficult Customers: HEARD and LAST, perceived justice, establishing common "
    "ground, setting limits and managing your own emotions (K1, K3, A3)",
 8: "Topic 9 — Understanding When to Escalate: functional vs hierarchical escalation, trigger behaviours, "
    "the three-strike rule, documentation and POHA (K3, A3)",
}
ADMIN_TEXT = ("Welcome, trainer and learner introductions, ground rules, learning outcomes, course outline, "
              "Skills Framework alignment and mandatory SSG digital attendance (AM)")
RECAP_TEXT = ("Course summary, key takeaways, TRAQOM course feedback survey and the Briefing for Assessment. "
              "Assessment digital attendance")
ASSESS_TEXT = {
 "OQ": "Oral Questioning (OQ) — 3 open-ended questions covering K1, K2 and K3. Individual, summative, open book",
 "RP": "Role Play (RP) — 2 simulated customer interactions covering A1, A2 and A3. Individual, summative, open book",
}

# (minutes, kind, ref)
SEQ = [
 (20, "admin", "ADMIN"),
 (30, "topic", 0), (20, "activity", 0),
 (25, "topic", 1),
 (15, "break", "Tea break"),
 (15, "activity", 1),
 (30, "topic", 2), (20, "activity", 2),
 (25, "topic", 3), (15, "activity", 3),
 (15, "topic", 4),
 (60, "lunch", "Lunch break"),
 (15, "activity", 4),
 (20, "topic", 5), (15, "activity", 5),
 (25, "topic", 6), (20, "activity", 6),
 (15, "break", "Tea break"),
 (20, "topic", 7), (15, "activity", 7),
 (20, "topic", 8), (15, "activity", 8),
 (10, "recap", "RECAP"),
 (30, "assess", "OQ"), (30, "assess", "RP"),
]

def _hm(x):
    return f"{x // 60}:{x % 60:02d}"

def build_rows():
    t = 9 * 60
    out = []
    for mins, kind, ref in SEQ:
        if kind == "topic":
            text = TOPIC_TEXT[ref]
        elif kind == "activity":
            text = act(ref + 1)
        elif kind == "admin":
            text = ADMIN_TEXT
        elif kind == "recap":
            text = RECAP_TEXT
        elif kind == "assess":
            text = ASSESS_TEXT[ref]
        else:
            text = ref
        out.append((_hm(t), _hm(t + mins), mins, kind, text))
        t += mins
    return out

SCHEDULE = {1: (C.DAY_THEMES[1], build_rows())}

# ---------------------------------------------------------------- build
doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc, "LESSON PLAN", C.TITLE, C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc, [
    ("1.0", "1 March 2026",
     "Initial release — 9 topics across 2 learning units, aligned to TSC EPW-CEX-3034-1.1.", C.TRAINER),
    (C.VERSION.lstrip("v"), C.VERSION_DATE,
     "Major content revision. Nine topics deepened from current industry practice (Zendesk, Help Scout, "
     "Qualtrics, SurveyMonkey, Coursera, Tidio, Intercom, Indeed SG, Global Response). Replaced generic "
     "exercises with nine real-world Singapore case-study activities, each with scenario, discussion "
     "questions and trainer debrief. Added the activities folder with per-activity PDFs.", C.TRAINER),
])
prodoc.add_toc(doc)


def H(text, level=1):
    return doc.add_heading(text, level=level)


def set_cell(cell, text, bold=False, size=9.5, color=None, fill=None, align=None):
    cell.text = ""; p = cell.paragraphs[0]
    if align: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = "Arial"
    if color: r.font.color.rgb = color
    if fill: prodoc._shade_cell(cell, fill)


H("Course Information", 1)
info = [("Course Title", C.TITLE),
        ("WSQ Course Reference", C.COURSE_CODE),
        ("Training Provider", C.ORG + "  (" + C.UEN.replace("UEN: ", "UEN ") + ")"),
        ("Skills Framework TSC", f"{C.TSC_TITLE}  ·  {C.TSC_CODE}"),
        ("Duration", "1 day · 8 hours (7 hours classroom facilitation + 1 hour assessment)"),
        ("Daily Timing", "9:00 am – 6:00 pm (1-hour lunch; tea breaks counted within training time)"),
        ("Mode", "Instructor-led classroom: lecture, peer sharing, case study and role play"),
        ("Trainer-to-learner ratio", "Lecture and peer sharing 1:3 (min) to 1:20 (max); role play 1:1"),
        ("Trainer", C.TRAINER)]
t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
for k, v in info:
    c = t.add_row().cells
    set_cell(c[0], k, bold=True, size=10, fill=TOPIC_FILL)
    set_cell(c[1], v, size=10)

H("Learning Outcomes", 1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size = Pt(10.5)

H("Learning Units", 1)
lt = doc.add_table(rows=0, cols=5); lt.style = "Table Grid"
hdr = lt.add_row().cells
for i, htext in enumerate(["Learning Unit", "Learning Outcome", "Topics", "K / A mapped", "Facilitation"]):
    set_cell(hdr[i], htext, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
for lu in C.LEARNING_UNITS:
    c = lt.add_row().cells
    set_cell(c[0], f"LU{lu['num']}: {lu['title']}", bold=True, size=9.5, fill=TOPIC_FILL)
    set_cell(c[1], lu["lo"], size=9.5)
    set_cell(c[2], lu["topics"], size=9.5)
    set_cell(c[3], lu["ka"], size=9.5)
    set_cell(c[4], lu["hours"], size=9.5)

H("Assessment", 1)
for a in [C.ASSESSMENT["written"], C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, Learner Guide and approved materials only.",
          "Oral Clarification of up to 10 minutes may be conducted 1:1 to close minor performance gaps; "
          "this is not counted in the assessment duration.",
          C.ASSESSMENT["note"]]:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size = Pt(10.5)

KIND_FILL = {"topic": TOPIC_FILL, "break": BREAK_FILL, "lunch": LUNCH_FILL,
             "assess": ASSESS_FILL, "admin": "F3F5F8", "recap": "F3F5F8", "activity": None}

H("Course Schedule", 1)
for day, (theme, rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}", 2)
    tbl = doc.add_table(rows=0, cols=3); tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.add_row().cells
    for i, htext in enumerate(["Time", "Duration", "Topic / Activity"]):
        set_cell(hdr[i], htext, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
    facilitation = 0; assessment = 0
    for start, end, mins, kind, text in rows:
        cells = tbl.add_row().cells; fill = KIND_FILL.get(kind)
        set_cell(cells[0], f"{start}–{end}", bold=(kind in ("topic", "assess")), size=9.5, fill=fill)
        set_cell(cells[1], f"{mins} min", size=9.5, fill=fill)
        set_cell(cells[2], text, bold=(kind in ("topic", "assess")), size=9.5, fill=fill)
        if kind == "assess":
            assessment += mins
        elif kind != "lunch":
            facilitation += mins
    for row in tbl.rows:
        row.cells[0].width = Inches(1.0); row.cells[1].width = Inches(0.8)
        row.cells[2].width = Inches(5.0)
    p = doc.add_paragraph()
    r = p.add_run(f"Classroom facilitation: {facilitation} minutes ({facilitation/60:.0f} hours).  "
                  f"Assessment: {assessment} minutes ({assessment/60:.0f} hour).  "
                  f"Total: {(facilitation+assessment)/60:.0f} hours excluding the 1-hour lunch break.")
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY
    assert facilitation == 420, f"facilitation = {facilitation}, expected 420 (7 h)"
    assert assessment == 60, f"assessment = {assessment}, expected 60 (1 h)"

H("Activity Reference", 1)
tt = doc.add_table(rows=0, cols=5); tt.style = "Table Grid"
hdr = tt.add_row().cells
for i, htext in enumerate(["Topic", "Activity", "Format", "Duration", "K / A mapped"]):
    set_cell(hdr[i], htext, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
for tp in C.TOPICS:
    acts = [a for a in ACT if a["topic"] == tp["num"]]
    for a in acts:
        cells = tt.add_row().cells
        set_cell(cells[0], f"T{tp['num']}: {tp['title']}", bold=True, size=9, fill=TOPIC_FILL)
        set_cell(cells[1], f"Activity {a['num']}: {a['title']}", size=9)
        set_cell(cells[2], a["type"], size=9)
        set_cell(cells[3], f"{a['minutes']} min", size=9)
        set_cell(cells[4], a["objective"].split("·")[1].strip() if "·" in a["objective"] else "", size=9)

H("Instructional Methods", 1)
im = doc.add_table(rows=0, cols=3); im.style = "Table Grid"
hdr = im.add_row().cells
for i, htext in enumerate(["Method", "Ratio", "How it is used in this course"]):
    set_cell(hdr[i], htext, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
for m, ratio, use in [
    ("Lecture", "1:3 – 1:20",
     "Trainer-led delivery of each topic's concepts using the slide deck, supported by real Singapore "
     "service examples, frameworks and comparisons."),
    ("Peer Sharing", "1:3 – 1:20",
     "Groups of 3–5 analyse a case study and share findings with the class. Used in Activities 1, 2, 3, 6 "
     "and 7, with a group presentation in Activity 7."),
    ("Role Play", "1:1",
     "Learners play the service officer and the customer in a simulated interaction, observed and debriefed. "
     "Used in Activities 2, 3, 4, 5, 8 and 9 — the same format as the Role Play assessment."),
]:
    c = im.add_row().cells
    set_cell(c[0], m, bold=True, size=9.5, fill=TOPIC_FILL)
    set_cell(c[1], ratio, size=9.5)
    set_cell(c[2], use, size=9.5)

H("Resources Required", 1)
for r_ in ["Classroom for up to 20 learners with movable seating for pair and group work.",
           "Projector / 75-inch display, whiteboard, markers and flip-chart paper for group mapping.",
           "Printed activity briefs from the activities folder (one per learner per activity).",
           "Course slide deck and Learner Guide, downloadable from the LMS at https://lms-tms.tertiaryinfotech.com/.",
           "SSG digital attendance QR (TRAQOM) for AM, PM and assessment attendance."]:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(r_).font.size = Pt(10.5)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
os.makedirs(OUTDIR, exist_ok=True)
OUT = os.path.join(OUTDIR, f"LP-{C.SHORT_TITLE}.docx")
doc.save(OUT)
print("Saved", OUT)
