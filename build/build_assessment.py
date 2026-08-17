#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""WSQ assessment set — Improve Your Business with Excellent Customer Service (TGS-2025056191).

MIRRORS THE ORIGINAL PAPER held on Drive (the ATO's filed version):
  - Oral Questioning (OQ) — 3 open-ended KNOWLEDGE questions, K1 / K2 / K3, 30 minutes
  - Role Play (RP)        — 2 simulated customer interactions, (A1, A2) / A3, 30 minutes
Same instrument names, same question counts, same K/A mapping and the same timings.
Only the scenarios, questions and model answers are rewritten, from THIS course's
slides, Learner Guide and activities.

Each instrument is produced as a Question Paper and a matching Answer Key (4 DOCX),
all with the WSQ house cover page. Page 1 cover; page 2 Trainee Information +
Instructions + Grading; the questions begin on page 3. Body Arial 11.
DOCX only — no PDFs for the assessment set.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0, HERE)
import course_data as C
import prodoc

REPO = os.path.dirname(HERE)                       # .../courseware
OUT = os.path.join(REPO, "assessment")
ASSETS = os.path.join(REPO, "assets")
os.makedirs(OUT, exist_ok=True)

TITLE = C.TITLE
COURSE_CODE = C.COURSE_CODE
prodoc.TGS = f"TGS Ref No: {COURSE_CODE}"
ORG_LOGO = os.path.join(ASSETS, "tertiary-infotech-logo.png")
COURSE_LOGO = None

Q_VER = A_VER = C.VERSION.lstrip("v")
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27)
GREY = RGBColor(0x55, 0x5B, 0x66)

# ================================================================ ORAL QUESTIONING
# (criterion, context, question, [model-answer points]) — each traces to the slides / LG.
ORAL = [
 ("K1",
  "You are a customer service representative at a retail electronics store in Singapore. A regular "
  "customer visits and mentions that service quality has felt inconsistent across their recent visits. "
  "Your manager asks you to explain how you would gather and understand this customer's concerns.",
  "Explain the key principles of effective communication you would apply when interacting with this "
  "customer to understand their concerns and collect meaningful feedback. Provide at least three "
  "principles and describe how each one helps in a customer service context.",
  ["Award the mark where the candidate names at least THREE principles and explains the service value of each. "
   "The list below is indicative, not exhaustive.",
   "ACTIVE LISTENING — attend fully, let the customer finish, and do not interrupt. In service this surfaces "
   "the real issue: the detail that matters is often buried in the first 30–45 seconds of an unstructured account.",
   "PARAPHRASING AND CONFIRMING — restate the concern in your own words and ask 'have I got that right?'. "
   "This makes listening visible, proves the customer was heard, and lets them correct you before you act on a "
   "wrong understanding.",
   "OPEN VS CLOSED QUESTIONING — open questions ('walk me through what happened') explore and surface the need; "
   "closed questions ('was it the Tampines branch?') confirm facts. Diagnose with open questions BEFORE proposing "
   "a solution; reversing the order turns the conversation into an interrogation.",
   "POSITIVE LANGUAGE — frame the response around what IS possible ('I can arrange an exchange today') rather "
   "than what is not ('we don't do refunds'). Same underlying answer, materially different customer experience.",
   "EMPATHY AND ACKNOWLEDGEMENT — name the customer's experience and its impact before moving to a fix. "
   "Acknowledging the experience is not the same as admitting fault.",
   "NON-VERBAL COMMUNICATION — open posture, roughly 70% eye contact, and a genuine greeting. A large share of "
   "the first impression is formed from appearance and body language, before any words are exchanged.",
   "CLARITY AND PLAIN LANGUAGE — avoid internal jargon, policy-section references and acronyms the customer "
   "cannot be expected to know.",
   "CLOSING WITH NEXT STEPS — state what will happen, by when, and who owns it, so the customer leaves with "
   "certainty rather than hope.",
   "A strong answer connects each principle to WHAT IT ACHIEVES in the service interaction, rather than simply "
   "listing the principles by name."]),

 ("K2",
  "You are a team leader at a Singapore hospitality company operating a hotel and a restaurant. Management "
  "has noticed a decline in online review scores over the past quarter and wants a more structured approach "
  "to collecting customer feedback across all touchpoints.",
  "Describe at least four different customer feedback channels your company could use to collect customer "
  "feedback. For each channel, explain one advantage and one limitation.",
  ["Award the mark where the candidate names at least FOUR channels, each with one genuine advantage and one "
   "genuine limitation. The pairs below are indicative.",
   "POST-VISIT SURVEYS (CSAT / NPS / CES) — Advantage: reaches many customers, gives a trackable number over "
   "time, and can be tied to a specific touchpoint. Limitation: low response rates create severe non-response "
   "bias, and a score alone rarely explains WHY.",
   "ONLINE REVIEWS (Google, TripAdvisor, social media) — Advantage: unprompted, candid, and visible to future "
   "customers, so it also shows reputational impact. Limitation: skews to the extremes — the delighted and the "
   "furious — and the silent middle is missing.",
   "SUPPORT TICKETS AND COMPLAINT LOGS — Advantage: captures specific, verifiable failures with dates and "
   "details, and can be counted and categorised. Limitation: only records customers who bothered to complain; "
   "roughly 56% of unhappy customers never say anything at all.",
   "DIRECT INTERVIEWS OR FEEDBACK CONVERSATIONS — Advantage: greatest diagnostic depth; you can ask follow-up "
   "questions and reach the root cause. Limitation: small sample, time-consuming, and subject to interviewer bias.",
   "FOCUS GROUPS — Advantage: surfaces shared themes and lets customers build on each other's points. "
   "Limitation: dominant participants can skew the discussion and it is not statistically representative.",
   "SOCIAL LISTENING — Advantage: captures unsolicited sentiment the customer never sent to you. "
   "Limitation: noisy, hard to attribute to a specific visit, and easily misread without context.",
   "COMMENT CARDS / QR FEEDBACK AT THE TOUCHPOINT — Advantage: immediate, captured while the experience is "
   "fresh. Limitation: very short responses with little diagnostic detail.",
   "FRONTLINE STAFF REPORTS — Advantage: staff see the causes that customers only experience as symptoms. "
   "Limitation: needs a formal channel, or the insight stays in informal chats and never reaches management.",
   "A strong answer notes that no single channel does both reach AND depth, so channels must be TRIANGULATED — "
   "using one alone risks fixing the wrong thing."]),

 ("K3",
  "You work in the operations department of a Singapore logistics company. Several customers have complained "
  "about delayed deliveries and poor communication from delivery personnel. Your supervisor has asked you to "
  "look into how internal feedback channels can be used to identify the root causes of these complaints and "
  "drive improvements.",
  "Explain what operation and process personnel feedback channels are, and describe at least three specific "
  "channels that can be used to gather feedback from frontline staff. How do these internal channels help in "
  "translating customer complaints into service improvements?",
  ["Award the mark where the candidate DEFINES the term, names at least THREE internal channels, and explains "
   "the link from complaint to improvement.",
   "DEFINITION — operation and process personnel feedback channels are the structured internal routes through "
   "which frontline and operational staff report what they observe about the service process: what is failing, "
   "why it is failing, and what would fix it. They gather feedback from STAFF about the PROCESS, as distinct "
   "from customer feedback channels which gather feedback from customers about their experience.",
   "SHIFT HUDDLES AND TEAM BRIEFINGS — short, regular, structured stand-ups where staff raise what went wrong "
   "on the last shift. The fastest route from observation to action.",
   "FRONTLINE INCIDENT LOGS — a written record of what actually happened at the point of failure, including "
   "detail the system never captured.",
   "TEAM RETROSPECTIVES / POST-INCIDENT REVIEWS — a periodic structured review of failures, with an owner "
   "assigned to each corrective action.",
   "INTERNAL ESCALATION REPORTS — the pattern in what gets escalated shows where frontline staff lack the "
   "authority or the information to resolve issues themselves.",
   "QUALITY ASSURANCE SCORECARDS — sampled interactions reviewed for tone, accuracy and process compliance, "
   "revealing systematic gaps rather than individual errors.",
   "SUGGESTION AND DEFECT-REPORTING CHANNELS — a formal route for staff to report a known system fault, so it "
   "does not sit unaddressed in an informal chat group.",
   "HOW THEY TRANSLATE COMPLAINTS INTO IMPROVEMENTS — customer channels report the SYMPTOM ('my delivery was "
   "late'); personnel channels usually supply the CAUSE ('the route planner does not account for the new "
   "one-way system, and drivers have no way to flag it'). Without the internal channel you can only treat "
   "symptoms case by case.",
   "THE CLOSED LOOP — collect (from both customers and staff) → analyse to separate symptom from root cause → "
   "act with a named owner, a date and a measure → tell the customer what changed. A strong answer states that "
   "the improvement must be specific and owned, not a general intention.",
   "A strong answer explicitly contrasts internal (staff/process) channels with external (customer) channels, "
   "and explains why BOTH are needed."]),
]

# ================================================================ ROLE PLAY
# (criterion, brief, task, [model-answer / marking-guide points])
ROLE_PLAY = [
 ("A1, A2",
  "You are a customer service representative at a fitness centre in Singapore. A new member, Mr Tan, "
  "approaches the front desk after his first week of membership. He looks slightly unsure and hesitant. "
  "Your task is to engage Mr Tan in conversation to collect his feedback about his experience so far, and "
  "to determine his needs and expectations regarding the centre's services and facilities. "
  "The assessor will play the role of Mr Tan.",
  "In this role play, demonstrate how you would greet Mr Tan, build rapport, use effective communication "
  "techniques to collect his feedback about his first-week experience, and identify his specific needs and "
  "expectations. You should use at least two different feedback collection methods during the interaction, "
  "and clearly determine what services or improvements would best meet his needs.",
  ["Mark the candidate COMPETENT where the following behaviours are observed. Wording will vary; look for the "
   "behaviour, not a script.",
   "GREETING AND RAPPORT — greets Mr Tan warmly within seconds, uses his name, open posture and genuine eye "
   "contact. Acknowledges that he is new and that it is his first week.",
   "COLLECTING FEEDBACK (A1) — uses at least TWO distinct feedback collection methods in the interaction. "
   "Acceptable combinations include: a structured verbal feedback conversation using open questions; offering "
   "a short post-onboarding survey or feedback form; a QR/app feedback prompt; scheduling a follow-up check-in "
   "call; or a comment card. The candidate must actually USE or OFFER both, not merely mention that they exist.",
   "OPEN QUESTIONING — explores with open questions: 'how has the first week gone?', 'walk me through a typical "
   "visit', 'what were you hoping to get out of joining?'. Closed questions used only to confirm specifics.",
   "ACTIVE LISTENING AND PARAPHRASING — lets Mr Tan finish, then restates his feedback in their own words and "
   "confirms it ('so the classes are good but the peak-hour queue for the machines is the frustration — have I "
   "got that right?').",
   "DETERMINING NEEDS AND EXPECTATIONS (A2) — separates what Mr Tan SAYS from what he NEEDS, and identifies "
   "both FUNCTIONAL needs (equipment availability, class timings, induction on machines) and EMOTIONAL needs "
   "(feeling welcome, not feeling like a beginner in front of experienced members). Identifying only functional "
   "needs is a partial answer.",
   "MATCHING SERVICES TO NEEDS — proposes specific services or improvements that address the needs identified: "
   "a complimentary induction session, off-peak timings, a personal trainer consultation, a class booking "
   "system, or a buddy/newcomer programme. The proposal must follow from what Mr Tan actually said.",
   "CONFIRMING AND CLOSING — summarises the needs back to Mr Tan, states concrete next steps with a time and an "
   "owner, and thanks him for the feedback.",
   "NOT YET COMPETENT indicators: proposing a solution before diagnosing; using only one feedback method; "
   "only closed questions; no paraphrasing; identifying no emotional need; or closing with no next step."]),

 ("A3",
  "You are a customer service officer at an online shopping platform's physical service centre in Singapore. "
  "Mrs Lim, a long-time customer, arrives visibly upset. She ordered a birthday gift for her daughter that "
  "was due three days ago and has not been delivered. She has already called the hotline twice with no "
  "resolution, and received an automated email saying the item is 'in transit' with no further detail. "
  "She is frustrated and is considering switching to a competitor. The assessor will play the role of Mrs Lim.",
  "In this role play, demonstrate how you would handle Mrs Lim's complaint, de-escalate her frustration, "
  "identify areas of service improvement based on her feedback, and recommend actionable steps to resolve her "
  "issue and prevent similar problems in future. Show appropriate use of emotional intelligence, escalation "
  "awareness and service recovery techniques.",
  ["Mark the candidate COMPETENT where the following behaviours are observed.",
   "DE-ESCALATION — lets Mrs Lim state her complaint fully without interrupting; stays calm in tone, pace and "
   "volume; acknowledges her frustration explicitly before addressing any facts.",
   "SERVICE RECOVERY FRAMEWORK — applies HEARD (Hear · Empathize · Apologize · Resolve · Diagnose) or LAST "
   "(Listen · Apologize · Solve · Thank). The framework need not be named, but the stages must be visible.",
   "EMPATHY WITH SPECIFICS — names the actual loss, not a generic phrase: the daughter's birthday, the three "
   "days, the two unanswered calls. 'I understand your frustration' on its own is insufficient.",
   "UNCONDITIONAL APOLOGY — apologises without 'if' or 'but', and without blaming a colleague, the courier or "
   "the system. Acknowledging the customer's experience is not the same as admitting legal fault.",
   "RESOLUTION WITH OWNERSHIP — offers a concrete resolution with a specific time and a named owner: trace the "
   "parcel now, arrange a replacement or refund, expedite delivery, or offer a goodwill gesture within their "
   "authority. Vague promises ('we'll look into it') do not meet the standard.",
   "ESCALATION AWARENESS — recognises the limits of their own authority and states clearly when and to whom "
   "they would escalate (a supervisor for a goodwill exception beyond their limit; the logistics team as a "
   "functional escalation to trace the parcel). Distinguishes escalation for AUTHORITY from escalation for "
   "EXPERTISE. Note that Mrs Lim is angry but not abusive — escalating on behaviour grounds is NOT appropriate here.",
   "IDENTIFYING AREAS OF IMPROVEMENT (A3) — identifies the ROOT CAUSES behind the complaint, not just the late "
   "parcel: the hotline calls produced no resolution and no callback; the automated 'in transit' email carried "
   "no real information; there is no proactive alert when a delivery misses its promised date; and no ownership "
   "was assigned across two prior contacts.",
   "RECOMMENDING ACTIONABLE STEPS — recommends specific, ownable improvements: proactive delay notifications "
   "before the customer has to chase; a case-ownership rule so a repeat caller is not restarted; replacing the "
   "generic status email with real tracking detail; and a callback commitment on unresolved hotline calls. "
   "Each recommendation should be specific enough to assign to someone.",
   "CLOSING THE LOOP — commits to a personal follow-up confirming the outcome, and to telling Mrs Lim what "
   "changed as a result of her feedback.",
   "NOT YET COMPETENT indicators: interrupting or becoming defensive; a conditional apology; blaming the "
   "courier or a colleague; offering no concrete resolution; identifying the late parcel as the only problem "
   "with no root-cause analysis; recommending only vague improvements; or no follow-up commitment."]),
]

# ================================================================ helpers
def base_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return doc


def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p


def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)


def answer_box(doc, lines=None, height_pt=110):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    if lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None)
            b.paragraph_format.left_indent = Inches(0.15)
            b.paragraph_format.space_after = Pt(3)
            rr = b.add_run("•  " + ln); rr.font.size = Pt(10.5)
    else:
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr(); trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt * 20)))
        trh.set(qn('w:hRule'), 'atLeast'); trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


FILL_GAP = 6


def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0


BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]
LMS_URL = "https://lms-tms.tertiaryinfotech.com/"


def add_hyperlink(p, url, text):
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link


def instructions(doc, minutes_text, n_items, what):
    heading(doc, "Instructions to Candidate")
    items = [
        f"This is the {what}.",
        f"A total of {minutes_text} is given to complete this assessment.",
        f"There are {n_items} questions. You need to answer all the questions.",
        "This is an open-book assessment that must be completed individually.",
        "You need to get all answers correct to be assessed as Competent.",
        None,
    ] + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)


def grading(doc, what):
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0


def finish(doc, path):
    prodoc.add_page_numbers(doc); prodoc.enable_update_fields(doc)
    doc.save(path); print("  saved:", os.path.basename(path))


# ================================================================ builders
def build_oq(answers):
    doc = base_doc()
    kind = "Oral Questioning (OQ) — Answer Key" if answers else "Oral Questioning (OQ)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Oral Questioning (OQ)" if answers else "Oral Questioning (OQ)",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        candidate_block(doc)
        instructions(doc, "30 minutes", len(ORAL), "Oral Questioning (OQ)")
        grading(doc, "Candidate has answered all questions and demonstrated the underpinning knowledge "
                     "(K1, K2, K3) required for the course learning outcomes.")
        page_break(doc)
    para(doc, "Questions and Answers", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Answer all questions in your own words. Each question tests underpinning knowledge covered "
              "in the course slides and the Learner Guide.",
         size=10.5, italic=True, color=GREY, after=8)
    # Explicit pagination: one question per page on the paper AND in the key — the
    # contexts and model answers are long. Never rely on keepNext/cantSplit.
    for i, (crit, ctx, q, pts) in enumerate(ORAL, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None, height_pt=170)
        if i < len(ORAL):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to OQ - {TITLE} - v{suffix}.docx" if answers
            else f"OQ - {TITLE} - v{suffix}.docx")
    finish(doc, os.path.join(OUT, name))


def build_rp(answers):
    doc = base_doc()
    kind = "Role Play (RP) — Answer Key" if answers else "Role Play (RP)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Role Play (RP)" if answers else "Role Play (RP)",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        candidate_block(doc)
        instructions(doc, "30 minutes", len(ROLE_PLAY), "Role Play (RP)")
        grading(doc, "Candidate has completed both role plays and demonstrated the abilities "
                     "(A1, A2, A3) required for the course learning outcomes.")
        page_break(doc)
    else:
        para(doc, "Marking guide for the assessor. Award Competent where the candidate demonstrates the "
                  "behaviours listed; the wording candidates use will vary and need not match this guide.",
             size=10.5, italic=True, color=GREY, after=8)
    para(doc, "Questions and Answers", size=13, bold=True, color=BRAND, after=4)
    if not answers:
        para(doc, "The assessor will play the role of the customer. You will be assessed on the behaviours "
                  "you demonstrate during the interaction, not on a written answer.",
             size=10.5, italic=True, color=GREY, after=8)
    for i, (crit, brief, task, pts) in enumerate(ROLE_PLAY, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, brief, size=11, after=3)
        para(doc, f"{task}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None, height_pt=170)
        if i < len(ROLE_PLAY):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to RP - {TITLE} - v{suffix}.docx" if answers
            else f"RP - {TITLE} - v{suffix}.docx")
    finish(doc, os.path.join(OUT, name))


if __name__ == "__main__":
    print("Building WSQ assessment set…")
    build_oq(answers=False); build_oq(answers=True)
    build_rp(answers=False); build_rp(answers=True)
    ks = sorted({c for c, *_ in ORAL})
    as_ = sorted({x.strip() for c, *_ in ROLE_PLAY for x in c.split(",")})
    print(f"Done. OQ: {len(ORAL)} questions ({', '.join(ks)}) · "
          f"RP: {len(ROLE_PLAY)} role plays ({', '.join(as_)}).")
    assert ks == ["K1", "K2", "K3"], f"OQ must cover K1-K3, got {ks}"
    assert as_ == ["A1", "A2", "A3"], f"RP must cover A1-A3, got {as_}"
